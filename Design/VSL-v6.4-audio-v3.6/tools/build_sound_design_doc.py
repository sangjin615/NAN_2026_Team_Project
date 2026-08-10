#!/usr/bin/env python3
"""Build the V6.4 sound design and implementation guide as a polished DOCX.

Preset: compact_reference_guide.
Named overrides: Malgun Gothic for Korean glyph coverage; editorial-cover title
29 pt navy, kicker 10 pt gold, subtitle 14 pt blue-gray.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT.parent / "미지의경매장_V6.4_사운드디자인_통합사양서.docx"
SOUND = json.loads((ROOT / "sound" / "sound.json").read_text(encoding="utf-8"))
FLOW = json.loads((ROOT / "flow.json").read_text(encoding="utf-8"))

FONT = "맑은 고딕"
NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "B88732"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
INK = "202124"
TOTAL_DXA = 9360
TABLE_INDENT = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, bold=None, italic=None, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    assert sum(widths) == TOTAL_DXA, widths
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("tblW", "tblInd", "tblLayout"):
        old = tbl_pr.find(qn(f"w:{tag}"))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(TOTAL_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    base = max(existing or [0]) + 1
    existing_num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_base = max(existing_num_ids or [0]) + 1
    ids = {}
    abstracts = []
    nums = []
    for offset, (kind, fmt, text) in enumerate((("bullet", "bullet", "●"), ("decimal", "decimal", "%1."))):
        abstract_id = base + offset
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        level.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        level.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        level.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), FONT)
        r_fonts.set(qn("w:hAnsi"), FONT)
        r_fonts.set(qn("w:eastAsia"), FONT)
        r_pr.append(r_fonts)
        level.append(r_pr)
        abstract.append(level)
        abstracts.append(abstract)

        num_id = num_base + offset
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        nums.append(num)
        ids[kind] = num_id
    first_num = numbering.find(qn("w:num"))
    insert_at = numbering.index(first_num) if first_num is not None else len(numbering)
    for offset, abstract in enumerate(abstracts):
        numbering.insert(insert_at + offset, abstract)
    for num in nums:
        numbering.append(num)
    return ids


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, NUMS["bullet"])
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        set_run_font(p.add_run(text))
    return p


def add_step(doc, text):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, NUMS["decimal"])
    set_run_font(p.add_run(text))
    return p


def add_para(doc, text="", *, bold=False, italic=False, size=11, color=INK, align=None, before=0, after=6):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    set_run_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def add_callout(doc, label, text, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TOTAL_DXA])
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{label}  "), bold=True, color=DARK_BLUE)
    set_run_font(p.add_run(text), color=INK)
    add_para(doc, "", after=2)


def add_table(doc, headers, rows, widths, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), size=9.3, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(table, widths)
    add_para(doc, "", after=2)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

NUMS = add_numbering(doc)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run_font(header.add_run("미지의 경매장 V6.4  |  SOUND DESIGN & WEB IMPLEMENTATION"), size=8.5, bold=True, color=MUTED)
footer = section.footer.paragraphs[0]
add_page_number(footer)

doc.core_properties.title = "미지의 경매장 V6.4 사운드디자인 통합사양서"
doc.core_properties.subject = "BGM, SFX, ambience, HTML/Web runtime and AI music production"
doc.core_properties.author = "미지의 경매장 프로젝트"
doc.core_properties.keywords = "game audio, VSL, HTML, web, BGM, SFX"

# Cover — editorial_cover.
add_para(doc, "", after=72)
add_para(doc, "SOUND DESIGN REFERENCE", bold=True, size=10, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
add_para(doc, "미지의 경매장 V6.4", bold=True, size=29, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
add_para(doc, "사운드디자인 · VSL 계약 · HTML/Web 런타임 통합사양서", size=14, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
add_para(doc, "빈티지 태엽 장치의 촉감과 낮은 밀도의 케이퍼 재즈를\n게임의 판단 리듬으로 연결하는 제작 기준", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=72)
add_para(doc, "VERSION 6.4  /  2026.08.01", bold=True, size=10.5, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
add_para(doc, "대상: HTML/Web · 음성 대사 없음 · 프로토타입 음원 포함", italic=True, size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

doc.add_heading("1. 결론과 제작 범위", level=1)
add_callout(doc, "우선안", f"BGM {len(SOUND['bgm'])}곡은 AI/작곡 시안을 2~3테이크씩 만든 뒤 DAW에서 루프·스템·믹스를 마감한다. SFX와 앰비언스는 직접 폴리/상업용 라이브러리를 중심으로 제작하고 생성형 도구는 빈 질감 보완에 사용한다.", fill="FFF8E8")
add_para(doc, "이번 납품은 ‘사운드 문서만’이 아니라 사운드 계약, 씬별 연결, 실제 웹 런타임, 교체 가능한 프로토타입 WAV를 한 구조로 묶는다.")
for text in (
    f"VSL 계약: 씬 {len(FLOW['nodes'])}개, UI 상태 {len(FLOW['uiStates'])}개, 행동 {len(FLOW['actions'])}개.",
    f"사운드 계약: SFX/앰비언스 큐 {len(SOUND['sfx'])}개, 행동 매핑 {len(SOUND['actionSfxMap'])}/{len(FLOW['actions'])}.",
    f"BGM: 작곡 단위 {len(SOUND['bgm'])}곡, 런타임 납품 슬롯 {SOUND['files']['expectedCount']['bgmFiles']}개(인트로·루프·레이어 포함).",
    "웹 구현: 사용자 제스처 언락, 마스터/BGM/SFX 버스, 씬 크로스페이드, 도시곡 연속 재생, 경매·유물 레이어링, 실패 팝업 피드백.",
    "음성: 대사·보이스오버는 제외하고 군중 웅성임과 장소 앰비언스만 사용.",
):
    add_bullet(doc, text)

doc.add_heading("2. V6 정합성 기준", level=1)
add_para(doc, "사운드가 잘못된 게임 규칙을 강화하지 않도록 V6 통합기획서를 우선 기준으로 삼고, v6.3 내 잔존 V5 항목을 정리했다.")
add_table(doc, ["항목", "V6.4 확정"], [
    ("족보 판매", "계약 수락·기한 없이 6종 패턴이 성립할 때 즉시 판매"),
    ("정보 채널", "수요 동향 · 출품 목록 · 경쟁자 예산; 유물 정보 채널 없음"),
    ("담보 대출", "상회 3단계, 처분가 45%, 2일 만기, 상환 x1.90"),
    ("유물 경매", "12일 경제 밖 3라운드; 거물은 이름만 표시"),
    ("감정", "확정가가 아닌 오차가 적용된 품질 범위 공개"),
], [2700, 6660], font_size=9.6)

doc.add_heading("3. 사운드 톤 앤 매너", level=1)
doc.add_heading("3.1 효과음 언어", level=2)
add_para(doc, "고급 골동품점의 단단한 재료감에 태엽식 장치의 작은 운동을 더한다. 장난스럽거나 과장된 카툰음보다 ‘정밀한 낡음’을 목표로 한다.")
add_table(doc, ["축", "사용", "피함"], [
    ("기계", "미세 기어, 래칫, 태엽 감김, 황동 걸쇠", "공장 소음, 무거운 디젤 엔진"),
    ("거래", "동전, 두꺼운 지폐, 영수증 스탬프, 목재 서랍", "카지노 잭팟, 모바일 보상 폭발음"),
    ("유리/도자기", "얇은 공명, 정돈되는 파편, 짧은 잔향", "날카로운 파손음, 공포성 충격음"),
    ("UI", "가죽 버튼, 종이 넘김, 작은 금속 클릭", "플라스틱 클릭, 현대 디지털 비프"),
], [1500, 3930, 3930], font_size=9.2)

doc.add_heading("3.2 음악 언어", level=2)
for text in (
    "희망: 근대 유럽 항구도시, 스팀펑크 장치, 항해를 앞둔 낙관. 선율은 명확하되 UI 판단을 가리지 않는다.",
    "미스터리: 1960년대 케이퍼 재즈의 어휘를 차용하되 특정 작품·아티스트를 모사하지 않는다. 콘트라베이스와 브러시가 중심이고 색소폰은 드문 문장부호처럼만 등장한다.",
    "밀도: 대사 없는 게임이지만 텍스트 읽기와 가격 판단을 방해하지 않도록 중저역을 정리하고 리드 악기의 체류 시간을 짧게 둔다.",
    "루프: 마지막 4~8마디에 새 정보를 넣지 않고 첫 마디로 자연스럽게 복귀한다. 인트로는 별도 파일로 분리한다.",
):
    add_bullet(doc, text)

doc.add_page_break()
doc.add_heading(f"4. BGM {len(SOUND['bgm'])}곡 통합 명세", level=1)
bgm_rows = []
for item in SOUND["bgm"]:
    layer_ids = [layer["id"] for layer in (item.get("layers") or [])]
    structure = " + ".join((["intro"] if item.get("structure", {}).get("intro") else []) + (layer_ids or ["loop"]))
    bgm_rows.append((item["id"], "·".join(item.get("scenes", [])), item["character"], structure))
add_table(doc, ["ID", "적용", "방향", "구조"], bgm_rows, [1700, 2100, 4300, 1260], font_size=7.8)

for item in SOUND["bgm"]:
    doc.add_heading(f"{item['name']} · {item['id']}", level=2)
    add_para(doc, f"{item['bpm']} BPM · {item['key']} · " + " · ".join(item["instrumentation"]))
    add_para(doc, item["character"])
    add_callout(doc, "AI 초안 프롬프트", item.get("generationPrompt", ""), fill=LIGHT_GRAY)

doc.add_heading("5. 장소 앰비언스", level=1)
add_para(doc, "도시의 진행 단계와 각 거점에 고유 BGM을 두되, 앰비언스는 음악보다 최소 10 dB 낮게 둔다. 명확한 대사처럼 들리는 군중 음성은 사용하지 않는다.")
add_table(doc, ["장소", "앰비언스", "운용"], [
    ("로딩", "천천히 맞물리는 기어·벨트", "로딩 종료 시 300 ms 페이드아웃"),
    ("도시 거점", "먼 항구·목재 간판·가벼운 바람", "도시곡 위에 아주 얕게"),
    ("감정소", "확대경·천·유리 진열장", "감정 행동 때만 근접음"),
    ("술집", "낮은 웅성임·유리잔·벽난로", "개별 단어 식별 금지"),
    ("거래소", "종이·장부·주판", "족보 성립 SFX 공간 확보"),
    ("조합", "금고·체인·두꺼운 문", "대출 행동 때만 저역 강조"),
    ("일반 경매", "낮은 군중·옷깃·의자", "망치와 입찰음을 가리지 않음"),
    ("마감일", "얇은 시계 초침", "3·6·9일 거점에서만 선택적"),
], [1600, 4140, 3620], font_size=9.2)

doc.add_heading("6. SFX 설계와 우선순위", level=1)
add_para(doc, f"총 {len(SOUND['sfx'])}개 큐는 UI·거래·감정·경매·진행·유물·앰비언스 그룹으로 관리한다. 모든 행동에는 정확히 하나의 기본 큐가 연결되어 있고, 결과의 좋고 나쁨을 미리 누설하지 않는 것이 원칙이다.")
add_table(doc, ["우선", "큐", "소리의 핵심", "트리거"], [
    ("P0", "sfx-hanbo-complete", "유리 조각과 황동 톱니가 제자리에 맞물리는 상승 3음", "족보 판매 성립"),
    ("P0", "sfx-gavel-hit", "건조한 목재 망치 + 짧은 홀 잔향", "낙찰 확정"),
    ("P0", "sfx-appraise-reveal", "천을 걷고 얇은 유리가 울리는 중립 공개음", "감정 결과 공개"),
    ("P0", "sfx-day-advance", "태엽을 한 칸 감고 멈추는 래칫", "다음 날"),
    ("P0", "sfx-relic-acquire", "무거운 잠금 해제 + 금속/합창의 짧은 광택", "유물 획득"),
    ("P1", "sfx-info-buy", "동전·종이·작은 서랍", "정보 구매"),
    ("P1", "sfx-loan-take / repay", "금고와 체인의 열림/회수", "대출 실행·상환"),
    ("P1", "sfx-failure", "낮은 금속 걸쇠와 짧은 종이 멈춤", "실패 팝업"),
], [850, 2050, 4060, 2400], font_size=8.8)
add_callout(doc, "믹스 기준", "BGM 목표 -20 LUFS-I 전후, 앰비언스는 BGM보다 10~14 dB 낮게, 일반 SFX -18~-16 LUFS short-term, 망치·족보·유물 같은 강조 SFX는 -14 LUFS short-term 부근. 실제 최종값은 통합 플레이에서 귀로 조정한다.")

doc.add_heading("7. 씬별 사운드 연결표", level=1)
scene_labels = {
    "scene-title": "타이틀", "scene-continue": "이어하기", "scene-loading": "로딩",
    "scene-city": "도시 지도", "scene-office": "감정소", "scene-tavern": "술집",
    "scene-exchange": "거래소", "scene-guild": "조합", "scene-merchant": "상회",
    "scene-auction": "일반 경매", "scene-summary": "하루 결산", "scene-ending": "12일 판정",
    "scene-final": "유물 경매", "scene-result": "최종 결과", "scene-museum": "유물 전시관",
}
scene_rows = []
for scene in FLOW["nodes"]:
    sid = scene["id"]
    mapping = SOUND["sceneBgmMap"].get(sid, {})
    bgm = mapping.get("bgm", "-")
    layers = ", ".join(mapping.get("layers") or []) or "loop"
    special = "동일 BGM 유지"
    if sid == "scene-loading": special = "기어 앰비언스"
    elif sid == "scene-auction": special = "군중 앰비언스·입찰 강도 레이어"
    elif sid == "scene-final": special = "라운드별 L1→L3"
    elif sid == "scene-summary": special = "순익/손실 큐"
    elif sid == "scene-result": special = "성공/파산 큐"
    scene_rows.append((scene_labels.get(sid, sid), bgm, layers, special))
add_table(doc, ["씬", "BGM", "기본 레이어", "전환/조건"], scene_rows, [1800, 2400, 2000, 3160], font_size=8.6)

doc.add_heading("8. HTML/Web 런타임", level=1)
add_para(doc, "런타임은 Web Audio의 복잡한 그래프 대신 교체와 디버깅이 쉬운 HTMLAudio 기반으로 구성했다. 현재 WAV는 기능 검증용 프로토타입이며, 같은 파일명으로 최종 에셋을 덮어쓰면 코드 수정 없이 교체된다.")
for text in (
    "첫 pointerdown/keydown에서 오디오를 언락하고 현재 씬 BGM을 시작한다.",
    "sceneBgmMap을 기준으로 씬과 공개된 일차·결산 결과에 맞는 변주를 선택한다.",
    "BGM이 바뀔 때 1.8초 크로스페이드, 결말 전환은 0.6초로 짧게 처리한다.",
    "경매의 공개 호가·LOT 순서, 유물 라운드, 공개된 결산 결과에 따라 동위상 레이어를 켜고 끈다.",
    "마스터/BGM/SFX 값은 localStorage에 저장하고 모든 활성 트랙에 즉시 반영한다.",
    "SFX는 큐별 폴리포니 제한을 두어 연타 시 오래된 재생을 정리한다.",
):
    add_bullet(doc, text)
add_table(doc, ["파일", "역할"], [
    ("sound/sound.json", "큐·BGM·씬·행동 매핑의 단일 기준"),
    ("assets/runtime/audio/sound-runtime.js", "file://에서도 읽는 브라우저 설정"),
    ("assets/runtime/audio/audio-manager.js", "재생·볼륨·크로스페이드·레이어"),
    ("assets/runtime/audio/{bgm,sfx,ambience}", "교체 가능한 WAV 슬롯"),
    ("assets/runtime/audio/prototype-manifest.json", "프로토타입 생성 파라미터와 파일 목록"),
], [3200, 6160], font_size=9.4)

doc.add_heading("9. AI 음악 제작 방법 비교", level=1)
add_table(doc, ["방법", "적합도", "장점", "주의"], [
    ("AIVA Pro", "1순위 BGM", "MIDI/오디오 편집, 다양한 스타일, Pro는 전체 저작권 조건", "업로드 영향물은 권리 보유 자료만 사용"),
    ("Stable Audio 유료", "SFX·앰비언스 보조", "텍스트/오디오 조건화, 음악과 효과음 목표", "상업 이용은 유료 라이선스 확인"),
    ("Adobe Firefly Soundtrack", "영상 기반 무드 시안", "vibe/style/purpose/energy/tempo로 빠른 변형", "베타·영상 중심; 제품 내 조건을 생성 시점에 재확인"),
    ("Suno Pro/Premier", "아이디어 스케치", "빠른 곡 후보와 상업 이용 권한", "게임 루프·동위상 스템 제어가 핵심 요구와 덜 맞음"),
    ("작곡가/사운드 디자이너", "최종 품질", "정확한 루프·스템·저작권 체인·반복 수정", "비용과 일정 필요"),
], [1700, 1550, 3330, 2780], font_size=8.6)
add_callout(doc, "추천 조합", f"{len(SOUND['bgm'])}곡은 곡별 2~3개 시안을 만든 뒤 인간 편곡자가 DAW에서 루프와 {SOUND['files']['expectedCount']['bgmFiles']}개 납품 슬롯을 정리한다. 효과음은 직접 폴리/상업 라이브러리를 먼저 쓰고, 생성형 도구는 비어 있는 질감 보완에 사용한다.", fill="FFF8E8")
add_para(doc, "중요: 서비스가 부여하는 상업 이용권과 각 국가에서 인정되는 저작권 보호는 같은 개념이 아니다. 출시 전 최종 약관·계정 플랜·생성일을 법무 또는 배급 기준으로 다시 확인한다.", italic=True, size=9.6, color="7A5A00")

doc.add_heading("10. 제작 워크플로", level=1)
for text in (
    "각 BGM 프롬프트로 4~8개 시안을 생성하고, 선율보다 루프 안정성·밀도·세계관 재료감을 먼저 평가한다.",
    "선택 시안의 MIDI 또는 스템을 DAW로 가져와 멜로디 과밀·저역 충돌·색소폰 빈도를 정리한다.",
    "BGM-03·15·16 경매곡과 BGM-04 유물곡은 정확히 같은 길이의 동위상 L1/L2/L3 스템으로 출력한다.",
    "-6 dBFS 이상의 헤드룸을 유지한 WAV 마스터와 게임용 OGG/AAC 파생본을 만든다.",
    "프로토타입 폴더의 동일 파일명으로 교체하고 전체 12일 플레이에서 전환·반복 피로·SFX 마스킹을 점검한다.",
    "최종 승인본의 생성 서비스, 플랜, 생성일, 프롬프트, 원본 다운로드, 편집 프로젝트, 영수증, 약관 캡처를 함께 보관한다.",
):
    add_step(doc, text)

doc.add_heading("11. 납품 규격", level=1)
add_table(doc, ["구분", "규격"], [
    ("BGM 마스터", "48 kHz / 24-bit WAV, stereo, intro/loop 또는 동위상 layer 분리"),
    ("SFX", "48 kHz / 24-bit WAV, mono 우선, 시작 무음 10 ms 이하"),
    ("앰비언스", "48 kHz / 24-bit WAV, stereo, seamless 30~90 s"),
    ("피크", "true peak -1 dBTP 이하 권장; 강조음끼리 중첩 테스트"),
    ("메타데이터", "cue ID, 버전, 길이, 루프 포인트, 라이선스 출처, 제작자"),
    ("웹 파생본", "브라우저 호환 OGG/AAC 추가 가능; WAV 슬롯은 기준 보존"),
], [2500, 6860], font_size=9.5)

doc.add_heading("12. 검수 체크리스트", level=1)
for text in (
    "[ ] 모든 71개 행동이 sound.json의 기본 큐에 연결되어 있다.",
    "[ ] 도시 내부 이동에서 BGM-02가 처음부터 재시작되지 않는다.",
    "[ ] 일반 경매의 L2/L3가 L1과 위상·길이가 일치한다.",
    "[ ] 유물 경매 1→3라운드 레이어 상승이 클릭 없이 자동 반영된다.",
    "[ ] 마스터/BGM/SFX 0%가 각각 완전 음소거되고 저장 후 복원된다.",
    "[ ] 팝업 열림/닫힘과 실패 피드백이 중복 재생되지 않는다.",
    "[ ] 감정 공개음이 결과의 좋고 나쁨을 누설하지 않는다.",
    "[ ] 군중 앰비언스에서 식별 가능한 문장·언어가 들리지 않는다.",
    "[ ] 20분 이상 루프 청취에서 클릭·레벨 점프·피로가 없다.",
    "[ ] 각 최종 음원의 라이선스 증빙과 편집 원본이 보관되어 있다.",
):
    add_bullet(doc, text)

doc.add_heading("13. 출처 및 약관 확인일", level=1)
add_para(doc, "확인일: 2026-08-01. 아래 서비스의 정책은 변경될 수 있으므로 실제 생성·출시 시점에 재확인한다.", italic=True, size=9.5, color=MUTED)
sources = [
    ("AIVA EULA", "https://www.aiva.ai/legal/1"),
    ("AIVA overview / pricing", "https://www.aiva.ai/blog/aiva-ai-composer-overview"),
    ("Stable Audio FAQ", "https://stableaudio.com/faqs"),
    ("Stable Audio user guide", "https://stableaudio.com/user-guide/interface"),
    ("Adobe Firefly Generate Soundtrack", "https://helpx.adobe.com/firefly/web/work-with-audio-and-video/work-with-audio/generate-soundtrack-for-an-uploaded-video.html"),
    ("Adobe Firefly FAQ", "https://helpx.adobe.com/sg/firefly/web/get-started/learn-the-basics/adobe-firefly-faq.html"),
    ("Suno ownership", "https://help.suno.com/en/articles/2416769"),
    ("Suno paid rights", "https://help.suno.com/en/articles/9601665"),
]
for label, url in sources:
    add_bullet(doc, f"{label}: {url}")

doc.add_page_break()
doc.add_heading("부록 A. 현재 프로토타입 상태", level=1)
add_para(doc, "프로젝트에는 최종 음악이 아니라 코드와 전환 검증을 위한 합성 프로토타입 WAV가 들어 있다. 최종 음원으로 오인하지 않도록 manifest에 prototype=true를 기록했다.")
add_table(doc, ["검증", "결과"], [
    ("VSL 계약 검사", "PASS · 데이터/행동/핀/등록소 위반 0"),
    ("사운드 계약 검사", "PASS · 행동 71/71 · 오류 0 · 주의 0"),
    ("JavaScript 구문", "PASS · audio-manager, sound-runtime, index inline"),
    ("브라우저 스모크", "PASS · 타이틀→도시→술집→경매, 3버스 볼륨, 리소스 오류 0"),
], [2600, 6760], font_size=9.5)

doc.save(OUT)
print(OUT)
